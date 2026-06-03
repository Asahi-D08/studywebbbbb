"""Build a Word document containing the IB Economics homework answer."""

from pathlib import Path

from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

ASSETS = Path("/Users/_key/.cursor/projects/Users-key-code-studyweb/assets")
DIAGRAM_1 = ASSETS / "diagram1-indirect-tax-cigarettes-final.png"
DIAGRAM_2 = ASSETS / "diagram2-subsidy-nicotine-patches-final.png"
OUTPUT = Path("/Users/_key/code/studyweb/IB_Econ_Homework_23M_1_SL_TZ2_1.docx")

PART_A_PARA_1 = (
    "Indirect tax is the first policy I will use. Cigarettes is a classic example of a "
    "demerit good, which mean a good that consumers over-consume because they don't "
    "realise the full costs. When people smoke, third parties suffer \u2014 non-smokers "
    "breathe the smoke, and the public hospital have to pay later for the lung diseases. "
    "This is what economists call a negative externality of consumption. So the marginal "
    "private benefit (MPB), the benefit one smoker get from one more cigarette, is "
    "higher than the marginal social benefit (MSB), the benefit to whole society. In "
    "free market the equilibrium is at Qm, where MPB cross MPC, but the socially optimum "
    "quantity is at Q*, where MSB cross MSC, and the triangle between Qm and Q* is the "
    "welfare loss (over-consumption). A specific excise tax \u2014 government add fix "
    "amount per packet \u2014 raise the cost of production, so the supply curve (MPC) "
    "shift upward by the amount of the tax, becoming a new curve MPC + tax. After the "
    "tax, the price paid by consumer rise from Pm to Pc, quantity fall from Qm down to "
    "Q*, and the welfare loss is eliminate. The government also collect tax revenue, "
    "which is the box between Pc and the producer's price P_prod, from 0 to Q*; this "
    "revenue can be spend on healthcare or anti-smoking campaign. One weakness, "
    "cigarettes have inelastic price elasticity of demand (PED) because they are "
    "addictive, so the tax must be very big to really change consumer behaviour."
)

PART_A_PARA_2 = (
    "The second policy is a consumer subsidy on a close substitute. Here I take "
    "nicotine patches as the example, which is a merit good \u2014 a good consumer "
    "under-consume because they don't see the full benefit. When someone use patches "
    "instead of smoking, the society also gain (less second-hand smoke, lower public "
    "health cost), so this is a positive externality of consumption, and MSB is above "
    "the MPB. Free market equilibrium happen at Qm, where MPB cross MPC, but the "
    "socially optimum quantity is at Q*, where MSB cross MSC, and Q* is bigger than Qm "
    "because the patches are under-consumed; the welfare loss triangle is on the right "
    "side. A subsidy is money the government give to producers to lower their cost, so "
    "the supply curve (MPC) shift downward by the amount of subsidy, becoming a new "
    "curve MPC \u2212 subsidy. With the subsidy, the price for consumer fall to Pc, "
    "quantity rise from Qm up to Q*, and the welfare loss is removed. Because patches "
    "and cigarettes are substitutes in consumption, when patches become cheaper many "
    "consumers switch and the demand for cigarettes (the demerit good) shift left, so "
    "the demerit good market also move closer to its own socially optimum quantity. "
    "The main weakness is the high opportunity cost \u2014 the subsidy money come from "
    "the government budget that could be use for schools or hospitals \u2014 and also "
    "the substitute must really be a close substitute, otherwise consumer will not "
    "switch at all."
)

PART_B_PARA_1 = (
    "A public good is a good that is non-rivalrous (one person using it does not reduce "
    "the amount available for others) and non-excludable (you cannot stop people who "
    "don't pay from using it). Examples are national defense, street lighting and flood "
    "defenses. The free market under-provide them because of the free rider problem "
    "\u2014 since people can use without paying, private firms cannot earn profit and "
    "so they don't produce, which create a missing market, a form of market failure. "
    "The question is whether the government should always step in and provide directly."
)

PART_B_PARA_2 = (
    "On one hand, direct provision have real advantages. The government can collect tax "
    "from everybody, which solve the free rider problem \u2014 for example the UK "
    "government finance the Royal Air Force through tax, and every citizen benefit "
    "whether they pay or not. It also support equity, because every person have access "
    "regardless of income; Singapore's NParks public parks are free for all, but a "
    "private park would charge entry fee that poor families maybe cannot afford. Some "
    "goods are also too sensitive to give to private firms \u2014 almost no country let "
    "private companies run the army for security reasons, for example the US military "
    "is directly provided by the federal government even though USA is very "
    "market-based."
)

PART_B_PARA_3 = (
    "On the other hand, direct provision often suffer from government failure, which "
    "happen when the intervention create a worse outcome than the original problem. "
    "Without profit motive, public providers can be inefficient and bureaucratic; "
    "critics say the UK NHS have very long waiting time because there is no "
    "competition. There is also high opportunity cost, since the money could be use "
    "for schools or hospitals. Public-private partnerships (PPP) often work better, "
    "where the government finance and a private firm produce \u2014 Hong Kong's MTR is "
    "partly listed on the stock market but with the government as main shareholder, "
    "and it is one of the most efficient metro system in the world. Many goods are "
    "also actually quasi-public goods, meaning they only partly fit the definition. A "
    "road is non-rivalrous when empty but rivalrous in traffic, and with electronic "
    "tolls it become excludable, which is why in France private companies like Vinci "
    "and APRR operate the autoroutes. Ronald Coase also showed that English "
    "lighthouses were historically built and financed by private firms through port "
    "fees."
)

PART_B_PARA_4 = (
    "In my judgment, public goods do not always need to be provided directly. The key "
    "distinction is between provision (production) and finance (paying for it) \u2014 "
    "the government can finance and let private firms produce efficiently, as in "
    "Sweden's school voucher system where the government fund education but private "
    "schools (friskolor) deliver it. However, for pure public goods of national "
    "interest like national defense and the judicial system, direct provision is still "
    "the best because the risk of private abuse is too high. So the answer depend on "
    "the type of good (pure or quasi-public), the capacity of the government "
    "(efficient like Singapore vs weak with corruption), and how important equity is. "
    "A blanket \u201calways\u201d statement is too simplistic \u2014 a mixed approach "
    "work better in the real world."
)


def add_body_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(11)


def add_diagram(doc: Document, image_path: Path, caption: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Cm(15))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        run.italic = True
        run.font.size = Pt(10)
    cap.paragraph_format.space_after = Pt(12)


def main() -> None:
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("Question 7 \u2014 23M.1.SL.TZ2.1")
    run.bold = True
    run.font.size = Pt(14)

    h_a = doc.add_paragraph()
    run = h_a.add_run(
        "Part (a) \u2014 Explain two types of government intervention that could be "
        "used to correct the market failure arising from the consumption of demerit "
        "goods. [10]"
    )
    run.bold = True
    run.font.size = Pt(12)
    h_a.paragraph_format.space_before = Pt(6)
    h_a.paragraph_format.space_after = Pt(8)

    add_body_paragraph(doc, PART_A_PARA_1)
    add_diagram(
        doc,
        DIAGRAM_1,
        "Diagram 1: Negative consumption externality (cigarettes) "
        "corrected by indirect tax",
    )

    add_body_paragraph(doc, PART_A_PARA_2)
    add_diagram(
        doc,
        DIAGRAM_2,
        "Diagram 2: Positive consumption externality (nicotine patches) "
        "corrected by consumer subsidy",
    )

    h_b = doc.add_paragraph()
    run = h_b.add_run(
        "Part (b) \u2014 Using real-world examples, discuss whether public goods "
        "should always be provided directly by the government. [15]"
    )
    run.bold = True
    run.font.size = Pt(12)
    h_b.paragraph_format.space_before = Pt(6)
    h_b.paragraph_format.space_after = Pt(8)

    add_body_paragraph(doc, PART_B_PARA_1)
    add_body_paragraph(doc, PART_B_PARA_2)
    add_body_paragraph(doc, PART_B_PARA_3)
    add_body_paragraph(doc, PART_B_PARA_4)

    doc.save(str(OUTPUT))
    print(f"Wrote: {OUTPUT}")


if __name__ == "__main__":
    main()
